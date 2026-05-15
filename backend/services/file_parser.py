"""파일 텍스트 추출 — PDF(+OCR) / DOCX / XLSX / TXT"""

import io
import re
from pathlib import Path

MAX_CHARS = 12000  # Claude context 효율을 위한 상한


def extract_text(filename: str, content: bytes) -> str:
    """파일 확장자에 따라 텍스트를 추출한다."""
    ext = Path(filename).suffix.lower()

    if ext in (".txt", ".csv", ".md"):
        return _decode(content)
    if ext == ".pdf":
        return _extract_pdf(content)
    if ext == ".docx":
        return _extract_docx(content)
    if ext in (".xlsx", ".xls"):
        return _extract_xlsx(content)
    return ""


def _decode(content: bytes) -> str:
    for enc in ("utf-8", "cp949", "euc-kr"):
        try:
            return content.decode(enc)[:MAX_CHARS]
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")[:MAX_CHARS]


def _extract_pdf(content: bytes) -> str:
    """PDF 텍스트 추출 — 텍스트 레이어 우선."""
    import fitz  # pymupdf

    doc = fitz.open(stream=content, filetype="pdf")
    text = ""
    has_image_pages = False
    for page in doc:
        page_text = page.get_text()
        if page_text.strip():
            text += page_text + "\n"
        else:
            has_image_pages = True
            text += f"[페이지 {page.number + 1}: 이미지 — Vision 분석 필요]\n"
        if len(text) >= MAX_CHARS:
            break
    doc.close()
    return text[:MAX_CHARS]


def extract_pdf_images(content: bytes, max_pages: int = 5) -> list[str]:
    """PDF 페이지를 base64 이미지로 변환 (Vision API용).

    Returns: list of base64-encoded PNG strings
    """
    import fitz
    import base64

    doc = fitz.open(stream=content, filetype="pdf")
    images = []
    for i, page in enumerate(doc):
        if i >= max_pages:
            break
        # 텍스트가 없는 페이지(이미지 PDF)만 변환하거나, 전체 변환
        pix = page.get_pixmap(dpi=200)
        png_bytes = pix.tobytes("png")
        b64 = base64.b64encode(png_bytes).decode("utf-8")
        images.append(b64)
    doc.close()
    return images


def _extract_docx(content: bytes) -> str:
    """DOCX 전체 텍스트 + 테이블 추출."""
    from docx import Document

    doc = Document(io.BytesIO(content))

    parts = []

    # 본문 텍스트
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    # 테이블 (금액, 단가 등이 테이블에 있는 경우가 많음)
    for i, table in enumerate(doc.tables):
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows_text.append(" | ".join(cells))
        if rows_text:
            parts.append(f"\n[표 {i+1}]")
            parts.extend(rows_text)

    text = "\n".join(parts)
    return text[:MAX_CHARS]


def _extract_xlsx(content: bytes) -> str:
    """XLSX 전체 시트 · 전체 행 추출 (빈 행 건너뜀, 금액 누락 방지)."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    text = f"시트목록: {', '.join(wb.sheetnames)}\n\n"

    for name in wb.sheetnames:
        ws = wb[name]
        text += f"[시트: {name}]\n"
        row_count = 0
        for row in ws.iter_rows(values_only=True):
            vals = [v for v in row if v is not None]
            if not vals:
                continue
            # 값을 문자열로 변환, 숫자는 포맷팅
            formatted = []
            for v in row:
                if v is None:
                    formatted.append("")
                elif isinstance(v, (int, float)):
                    formatted.append(str(v))
                else:
                    formatted.append(str(v).strip())
            # 빈 셀만 있는 건 스킵
            line = " | ".join(c for c in formatted if c)
            if line.strip():
                text += f"  R{row_count+1}: {line}\n"
            row_count += 1
            if len(text) >= MAX_CHARS:
                break
        text += "\n"
        if len(text) >= MAX_CHARS:
            break

    wb.close()
    return text[:MAX_CHARS]
